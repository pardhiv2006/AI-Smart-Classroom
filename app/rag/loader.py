"""
ScholarMind AI — Document Loader
Handles PDF (page-range selective), DOCX, and TXT parsing.
Extracted and cleaned from original app.py.
"""
import io
import os
import streamlit as st
import fitz  # PyMuPDF
from langchain_core.documents import Document


# ─────────────────────────────────────────────────────────────────
# PDF HELPERS
# ─────────────────────────────────────────────────────────────────

def inspect_pdf_pages(file_bytes: bytes) -> int:
    """Return total page count of a PDF. Returns 0 on error."""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        return len(doc)
    except Exception:
        return 0


def estimate_ram_mb(page_count: int, chunk_size: int = 500) -> float:
    """
    Rough RAM estimate for processing page_count pages.
    ~3KB text per page → chunks → ~200 bytes per embedding vector (384-dim float32).
    """
    avg_text_per_page_kb = 3
    avg_chunks_per_page = max(1, (avg_text_per_page_kb * 1024) // chunk_size)
    total_chunks = page_count * avg_chunks_per_page
    # each 384-dim float32 vector = 384 * 4 bytes ≈ 1.5KB, plus overhead
    vector_mb = (total_chunks * 1.5) / 1024
    text_mb = (page_count * avg_text_per_page_kb) / 1024
    return round(text_mb + vector_mb, 1)


def process_pdf_pages(file_bytes: bytes, page_range: tuple[int, int], file_name: str = "uploaded_book.pdf") -> list[Document]:
    """
    Extract text from selected PDF pages (0-indexed range).
    Only processes pages with actual text content (skips scanned images).
    Returns list of LangChain Document objects.
    """
    docs = []
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        start = max(0, page_range[0])
        end = min(len(doc) - 1, page_range[1])
        for idx in range(start, end + 1):
            text = doc.load_page(idx).get_text()
            if text.strip():
                docs.append(Document(
                    page_content=text,
                    metadata={"book_name": file_name, "source": file_name, "page": idx + 1}
                ))
    except Exception as e:
        st.error(f"❌ PDF read error: {e}")
    return docs


# ─────────────────────────────────────────────────────────────────
# NON-PDF HELPERS
# ─────────────────────────────────────────────────────────────────

def process_docx(file_bytes: bytes, file_name: str) -> list[Document]:
    """Parse a DOCX file into a single Document."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        if text:
            return [Document(page_content=text, metadata={"book_name": file_name, "source": file_name, "page": 1})]
    except Exception as e:
        st.error(f"❌ DOCX read error: {e}")
    return []


def process_txt(file_bytes: bytes, file_name: str) -> list[Document]:
    """Parse a TXT/MD file into a single Document."""
    try:
        text = file_bytes.decode("utf-8", errors="ignore")
        if text.strip():
            return [Document(page_content=text, metadata={"book_name": file_name, "source": file_name, "page": 1})]
    except Exception as e:
        st.error(f"❌ TXT read error: {e}")
    return []


def process_non_pdf(file_bytes: bytes, file_name: str) -> list[Document]:
    """Route non-PDF files to the correct parser."""
    ext = os.path.splitext(file_name)[1].lower()
    if ext == ".docx":
        return process_docx(file_bytes, file_name)
    elif ext in (".txt", ".md"):
        return process_txt(file_bytes, file_name)
    else:
        st.warning(f"⚠️ Unsupported file type: {ext}")
        return []


# ─────────────────────────────────────────────────────────────────
# UNIFIED ENTRY POINT
# ─────────────────────────────────────────────────────────────────

def load_documents(
    file_bytes: bytes,
    file_name: str,
    page_range: tuple[int, int] | None = None,
) -> list[Document]:
    """
    Load and return Document objects from any supported file.
    For PDFs: page_range = (start_page_1indexed, end_page_1indexed).
    """
    is_pdf = file_name.lower().endswith(".pdf")
    if is_pdf:
        if page_range is None:
            total = inspect_pdf_pages(file_bytes)
            page_range = (1, total)
        # Convert to 0-indexed
        return process_pdf_pages(file_bytes, (page_range[0] - 1, page_range[1] - 1), file_name)
    else:
        return process_non_pdf(file_bytes, file_name)
